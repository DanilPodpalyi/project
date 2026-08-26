"""Тесты чтения исходных документов и схемы вакансии."""

import tempfile
import unittest
from pathlib import Path

from docx import Document

from extract_vacancy import VacancyData, read_vacancy


class VacancyParserTests(unittest.TestCase):
    def test_all_schema_fields_are_required_for_model_output(self) -> None:
        schema = VacancyData.model_json_schema()

        self.assertEqual(set(schema["required"]), set(schema["properties"]))

    def test_reads_vacancy_from_docx(self) -> None:
        with tempfile.TemporaryDirectory() as directory:
            file_path = Path(directory) / "vacancy.docx"
            document = Document()
            document.add_paragraph("PR-менеджер")
            document.add_paragraph("Требования: деловые коммуникации")
            document.save(file_path)

            self.assertEqual(
                read_vacancy(file_path),
                "PR-менеджер\nТребования: деловые коммуникации",
            )

    def test_reads_vacancy_from_txt_with_bom(self) -> None:
        with tempfile.TemporaryDirectory() as directory:
            file_path = Path(directory) / "vacancy.txt"
            file_path.write_text("PR-менеджер", encoding="utf-8-sig")

            self.assertEqual(read_vacancy(file_path), "PR-менеджер")


if __name__ == "__main__":
    unittest.main()
