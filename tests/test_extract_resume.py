"""Тесты чтения TXT, DOCX и PDF и схемы резюме."""

import tempfile
import unittest
from pathlib import Path

from docx import Document
from pypdf import PdfWriter
from pypdf.generic import DecodedStreamObject, DictionaryObject, NameObject

from extract_resume import ResumeData, _is_candidate_name_placeholder, read_resume


class ReadResumeTests(unittest.TestCase):
    def setUp(self) -> None:
        self.temporary_directory = tempfile.TemporaryDirectory()
        self.directory = Path(self.temporary_directory.name)

    def tearDown(self) -> None:
        self.temporary_directory.cleanup()

    def test_reads_utf8_txt_with_bom(self) -> None:
        file_path = self.directory / "resume.TXT"
        file_path.write_text("Имя: Иван", encoding="utf-8-sig")

        self.assertEqual(read_resume(file_path), "Имя: Иван")

    def test_reads_docx_paragraphs_and_tables_in_document_order(self) -> None:
        file_path = self.directory / "resume.docx"
        document = Document()
        document.add_paragraph("Имя: Иван")
        table = document.add_table(rows=1, cols=2)
        table.cell(0, 0).text = "Навык"
        table.cell(0, 1).text = "Python"
        document.add_paragraph("Опыт: 3 года")
        document.save(file_path)

        self.assertEqual(
            read_resume(file_path),
            "Имя: Иван\nНавык\tPython\nОпыт: 3 года",
        )

    def test_reads_pdf_text_layer(self) -> None:
        file_path = self.directory / "resume.pdf"
        writer = PdfWriter()
        page = writer.add_blank_page(width=612, height=792)
        font = DictionaryObject(
            {
                NameObject("/Type"): NameObject("/Font"),
                NameObject("/Subtype"): NameObject("/Type1"),
                NameObject("/BaseFont"): NameObject("/Helvetica"),
            }
        )
        font_reference = writer._add_object(font)
        page[NameObject("/Resources")] = DictionaryObject(
            {
                NameObject("/Font"): DictionaryObject(
                    {NameObject("/F1"): font_reference}
                )
            }
        )
        content = DecodedStreamObject()
        content.set_data(b"BT /F1 12 Tf 72 720 Td (Python Developer) Tj ET")
        page[NameObject("/Contents")] = writer._add_object(content)
        with file_path.open("wb") as output:
            writer.write(output)

        self.assertEqual(read_resume(file_path), "Python Developer")

    def test_reports_pdf_without_text_layer(self) -> None:
        file_path = self.directory / "scan.pdf"
        writer = PdfWriter()
        writer.add_blank_page(width=612, height=792)
        with file_path.open("wb") as output:
            writer.write(output)

        with self.assertRaisesRegex(SystemExit, "OCR"):
            read_resume(file_path)

    def test_reports_old_doc_format(self) -> None:
        file_path = self.directory / "resume.doc"
        file_path.write_bytes(b"legacy doc")

        with self.assertRaisesRegex(SystemExit, "сохранить как .docx"):
            read_resume(file_path)

    def test_reports_corrupted_docx(self) -> None:
        file_path = self.directory / "broken.docx"
        file_path.write_bytes(b"not a docx")

        with self.assertRaisesRegex(SystemExit, "Некорректный DOCX-файл"):
            read_resume(file_path)

    def test_all_schema_fields_are_required_for_model_output(self) -> None:
        schema = ResumeData.model_json_schema()

        self.assertEqual(set(schema["required"]), set(schema["properties"]))
        for definition in schema["$defs"].values():
            self.assertEqual(
                set(definition["required"]),
                set(definition["properties"]),
            )

    def test_recognizes_candidate_name_placeholders(self) -> None:
        self.assertTrue(_is_candidate_name_placeholder("ФИО"))
        self.assertTrue(
            _is_candidate_name_placeholder("Фамилия, имя, отчество")
        )
        self.assertFalse(_is_candidate_name_placeholder("Иван Петров"))


if __name__ == "__main__":
    unittest.main()
